import os
import streamlit as st
from app.modules.generate_doc import generate_resume, generate_cover_letter
from app.modules.cheat_sheet import generate_cheat_sheet, extract_topics_from_text, generate_combined_cheat_sheet
from app.modules.career_map import generate_career_map
from app.utils import fetch_jobs_remotive, fetch_jobs_adzuna, get_user_country
from app.modules.voice_interview.webrtc_audio import (
create_webrtc_recorder,
process_recorded_audio,
initialize_webrtc_session_state,
)
from app.modules.voice_interview import generate_feedback_from_audio
from app.modules.voice_interview.record_audio import save_uploaded_audio
initialize_webrtc_session_state()
st.set_page_config(page_title="AI Career Companion", layout="wide", page_icon="💼")
st.markdown("""
<style>
/* Global Styles */
* {
    font-family: 'Inter', 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
}

body {
    background: #0f172a;
    color: #e2e8f0;
}

/* Header Styles */
h1, h2, h3, h4, h5, h6 {
    color: #94a3b8;
    font-weight: 600;
}

/* Dark theme sidebar */
[data-testid="stSidebar"] {
    background: linear-gradient(135deg, #0f172a, #1e293b);
    border-right: 1px solid #334155;
    padding: 1.2rem 1rem;
    color: #cbd5e1;
    scrollbar-width: thin;
    scrollbar-color: #4a5a7a transparent;
}

/* Sidebar nav buttons */
.sidebar-nav-btn {
    display: block;
    width: 100%;
    padding: 1rem 1.5rem;
    margin-bottom: 0.75em;
    background: #1e293b;
    border: 1px solid #334155;
    border-radius: 8px;
    color: #cbd5e1;
    font-weight: 500;
    font-size: 1rem;
    text-align: left;
    transition: all 0.3s ease;
    cursor: pointer;
    user-select: none;
    position: relative;
    overflow: hidden;
}

.sidebar-nav-btn::before {
    content: '';
    position: absolute;
    top: 0;
    left: -100%;
    width: 100%;
    height: 100%;
    background: linear-gradient(90deg, transparent, rgba(37, 99, 235, 0.3), transparent);
    transition: 0.5s;
}

.sidebar-nav-btn:hover::before {
    left: 100%;
}

.sidebar-nav-btn.selected,
.sidebar-nav-btn:hover {
    background: rgba(30, 58, 138, 0.3);
    border-color: #3b82f6;
    color: #e0f2fe;
    box-shadow: 0 4px 12px rgba(37, 99, 235, 0.3);
}

/* Main content area */
section.main > div:first-child {
    background: rgba(15, 23, 42, 0.7);
    border: 1px solid #334155;
    border-radius: 12px;
    padding: 2rem;
    margin-top: 1.5rem;
    box-shadow: 0 4px 20px rgba(0, 0, 0, 0.4);
    backdrop-filter: blur(10px);
}

/* Input fields */
.stTextArea, .stTextInput, .stSelectbox, .stFileUploader {
    background: #1e293b !important;
    border: 1px solid #334155 !important;
    border-radius: 8px !important;
    color: #e2e8f0 !important;
}

.stTextArea textarea, .stTextInput input {
    color: #e2e8f0 !important;
}

/* Buttons */
.stButton>button {
    background: linear-gradient(135deg, #2563eb, #1d4ed8) !important;
    color: white !important;
    border: none !important;
    border-radius: 8px !important;
    padding: 0.5rem 1.5rem !important;
    font-weight: 500 !important;
    transition: all 0.3s ease !important;
    box-shadow: 0 2px 6px rgba(0, 0, 0, 0.2) !important;
}

.stButton>button:hover {
    transform: translateY(-2px) !important;
    box-shadow: 0 4px 10px rgba(37, 99, 235, 0.4) !important;
}

/* Tabs */
.stTabs [data-baseweb="tab-list"] {
    gap: 10px;
}

.stTabs [data-baseweb="tab"] {
    background: #1e293b !important;
    border-radius: 8px !important;
    padding: 0.5rem 1rem !important;
    margin: 0 5px !important;
    transition: all 0.3s ease !important;
}

.stTabs [aria-selected="true"] {
    background: linear-gradient(135deg, #2563eb, #1d4ed8) !important;
    color: white !important;
    box-shadow: 0 2px 6px rgba(37, 99, 235, 0.3) !important;
}

/* Cards */
.card {
    background: rgba(30, 41, 59, 0.7);
    border: 1px solid #334155;
    border-radius: 10px;
    padding: 1.5rem;
    margin-bottom: 1rem;
    transition: all 0.3s ease;
    box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);
}

.card:hover {
    transform: translateY(-3px);
    box-shadow: 0 6px 12px rgba(37, 99, 235, 0.2);
}

/* Footer */
footer {
    text-align: center;
    padding: 1rem;
    margin-top: 2rem;
    color: #94a3b8;
    font-size: 0.9rem;
}

/* Scrollbar styling */
::-webkit-scrollbar {
    width: 8px;
    height: 8px;
}

::-webkit-scrollbar-track {
    background: #0f172a;
}

::-webkit-scrollbar-thumb {
    background: #334155;
    border-radius: 4px;
}

::-webkit-scrollbar-thumb:hover {
    background: #475569;
}

/* Resume box styling */
.resume-box {
    background: rgba(30, 41, 59, 0.8);
    border-radius: 8px;
    padding: 24px 32px;
    border: 1px solid #334155;
    font-family: 'Segoe UI', Arial, sans-serif;
    font-size: 1.05rem;
    line-height: 1.7;
    margin-bottom: 1.5em;
    box-shadow: 0 4px 12px rgba(0, 0, 0, 0.2);
}

.resume-box ul { 
    margin-left: 1.5em;
    padding-left: 0;
}

.resume-box li { 
    margin-bottom: 0.5em; 
}

.resume-box h2, .resume-box h3, .resume-box h4 { 
    margin-top: 1.2em;
    color: #93c5fd;
    border-bottom: 1px solid #334155;
    padding-bottom: 0.3em;
}

/* Progress bar */
.stProgress > div > div > div {
    background: linear-gradient(90deg, #2563eb, #1d4ed8) !important;
}
</style>
""", unsafe_allow_html=True)

st.title("💼 AI Career Companion")

# --- Sidebar Navigation ---
nav_options = [
    ("🏠 Home", "home"),
    ("🤖 Automated Job Hunter", "job_hunter"),
    ("📄 Resume/Cover Letter", "resume"),
    ("📚 Cheat Sheet Generator", "cheatsheet"),
    ("🎤 Mock Interview", "interview")
    
]
if "nav_page" not in st.session_state:
    st.session_state["nav_page"] = "home"

st.sidebar.markdown("<div style='margin-bottom:1.5em;'><b style='font-size:1.2em;'>📂 Navigation</b></div>", unsafe_allow_html=True)
for label, key in nav_options:
    btn_class = "sidebar-nav-btn selected" if st.session_state["nav_page"] == key else "sidebar-nav-btn"
    if st.sidebar.button(label, key=f"navbtn_{key}", use_container_width=True):
        st.session_state["nav_page"] = key

# --- Main Dispatcher ---
if st.session_state["nav_page"] == "home":
    st.header("🏠 Welcome to AI Career Companion")
    st.markdown("""
    <div class='card'>
    <b>AI Career Companion</b> is your all-in-one assistant for job seekers! Upload your resume and a job description, and our app will:
    <ol>
        <li><b>Find relevant jobs</b> and show you matches</li>
        <li><b>Generate a tailored resume</b> for the job</li>
        <li><b>Create a custom cover letter</b></li>
        <li><b>Auto-generate a cheat sheet</b> of key topics</li>
        <li><b>Prepare you with a mock interview</b> based on your profile and the job</li>
    </ol>
    Everything is automated for you in one click!
    </div>
    """, unsafe_allow_html=True)
    with st.expander("🚀 Get Started: Upload Resume & Job Description", expanded=True):
        col1, col2 = st.columns(2)
        with col1:
            uploaded_resume = st.file_uploader("📄 Upload your resume (PDF)", type=["pdf"], key="home_resume")
        with col2:
            jd = st.text_area("📋 Paste Job Description", key='home_jd', height=150)
    if uploaded_resume and jd:
        from app.modules.generate_doc.resume_parser import extract_resume_details_from_pdf
        resume_details = extract_resume_details_from_pdf(uploaded_resume)
        resume_text = resume_details.get('raw_text', '')
        # Store in session state for all features
        st.session_state['shared_resume_details'] = resume_details
        st.session_state['shared_resume_text'] = resume_text
        st.session_state['shared_jd'] = jd
        st.success("Resume and Job Description uploaded! Ready to run full workflow.")
    run_workflow = st.button("▶️ Run Full AI Career Workflow", use_container_width=True, type="primary", disabled=not (uploaded_resume and jd))
    if run_workflow and uploaded_resume and jd:
        # 1. Job Hunter
        st.subheader("1️⃣ Automated Job Hunter")
        jobs = fetch_jobs_remotive(jd, limit=5, location='India')
        if not jobs:
            st.warning("No jobs found. Try different keywords or sources.")
        else:
            st.markdown(f"Found {len(jobs)} jobs matching your JD.")
            import re
            for i, job in enumerate(jobs, 1):
                job_desc_clean = re.sub(r'<[^>]+>', '', job['desc'])
                st.markdown(f"**{i}. [{job['title']}]({job['url']}) at {job['company']} ({job['location']})**")
                st.markdown(job_desc_clean[:300] + '...')
        # 2. Tailored Resume & Cover Letter
        st.subheader("2️⃣ Tailored Resume & Cover Letter")
        resume = generate_resume(jd, resume_details)
        st.markdown("<div class='resume-box'>" + resume.replace('\n', '<br>') + "</div>", unsafe_allow_html=True)
        from app.modules.generate_doc import genai
        import re
        name = ''
        for line in resume.split('\n'):
            if len(re.findall(r'[A-Z][a-z]+', line)) >= 2 and len(line.split()) <= 5:
                name = line.strip()
                break
        cover_letter_prompt = f"Write a professional cover letter for the following job description using this resume.\n\nJob Description:\n{jd}\n\nResume:\n{resume}\n\nMy name is {name}."
        model = genai.GenerativeModel('models/gemini-1.5-flash-latest')
        cl_response = model.generate_content(cover_letter_prompt)
        cl = cl_response.text if hasattr(cl_response, 'text') else str(cl_response)
        st.markdown("<div class='resume-box'>" + cl.replace('\n', '<br>') + "</div>", unsafe_allow_html=True)
        # 3. Cheat Sheet
        st.subheader("3️⃣ Cheat Sheet Generator")
        topics = extract_topics_from_text(resume_text, jd)
        st.markdown(f"Detected Topics: {', '.join(topics)}")
        for topic in topics:
            sheet = generate_cheat_sheet(topic, jd)
            with st.expander(f"📌 {topic}", expanded=False):
                st.markdown(sheet, unsafe_allow_html=True)
        # 4. Mock Interview
        st.subheader("4️⃣ Mock Interview")
        from app.modules.voice_interview import get_questions_from_resume_and_jd
        question_list = get_questions_from_resume_and_jd(resume_text, jd)
        if question_list:
            st.markdown(f"Generated {len(question_list)} interview questions:")
            for q in question_list:
                st.markdown(f"- {q['question']}")
        else:
            st.info("No interview questions generated.")

elif st.session_state["nav_page"] == "resume":
    st.header("📄 Resume & Cover Letter Generator")
    with st.expander("⚙️ Upload Resume & Job Description", expanded=True):
        col1, col2 = st.columns(2)
        with col1:
            uploaded_resume = st.file_uploader("📄 Upload your resume (PDF)", type=["pdf"])
        with col2:
            jd = st.text_area("📋 Paste Job Description", key='jd', height=150)
    # Use shared resume/JD if available and nothing uploaded
    resume_details = None
    if uploaded_resume:
        from app.modules.generate_doc.resume_parser import extract_resume_details_from_pdf
        resume_details = extract_resume_details_from_pdf(uploaded_resume)
        st.session_state['shared_resume_details'] = resume_details
        st.session_state['shared_resume_text'] = resume_details.get('raw_text', '')
    elif st.session_state.get('shared_resume_details'):
        resume_details = st.session_state['shared_resume_details']
        uploaded_resume = True  # For UI
    if not jd and st.session_state.get('shared_jd'):
        jd = st.session_state['shared_jd']
    with st.expander("📃 Extracted Resume Text"):
        st.text_area("Resume Text", resume_details.get('raw_text', '') if resume_details else '', height=200, label_visibility="collapsed")

    col1, col2 = st.columns(2)
    with col1:
        generate_resume_clicked = st.button("✨ Generate Resume", use_container_width=True)
    with col2:
        generate_cover_letter_clicked = st.button("📧 Generate Cover Letter", use_container_width=True)

    if generate_resume_clicked:
        with st.spinner("⏳ Generating your resume..."):
            resume = generate_resume(jd, resume_details)
            st.markdown("<div class='resume-box'>", unsafe_allow_html=True)
            st.markdown(resume.replace('\n', '<br>'), unsafe_allow_html=True)
            st.markdown("</div>", unsafe_allow_html=True)
            st.session_state['show_suggestions'] = True
            st.session_state['resume'] = resume

    if generate_cover_letter_clicked:
        improved_resume_text = st.session_state.get('improved_resume_text') or st.session_state.get('resume') or (resume_details.get('raw_text', '') if resume_details else '')
        import re
        name = ''
        for line in improved_resume_text.split('\n'):
            if len(re.findall(r'[A-Z][a-z]+', line)) >= 2 and len(line.split()) <= 5:
                name = line.strip()
                break
        cover_letter_prompt = f"Write a professional cover letter for the following job description using this resume.\n\nJob Description:\n{jd}\n\nResume:\n{improved_resume_text}\n\nMy name is {name}."
        from app.modules.generate_doc import genai
        model = genai.GenerativeModel('models/gemini-1.5-flash-latest')
        cl_response = model.generate_content(cover_letter_prompt)
        cl = cl_response.text if hasattr(cl_response, 'text') else str(cl_response)
        st.markdown("<div class='resume-box'>", unsafe_allow_html=True)
        st.markdown(cl.replace('\n', '<br>'), unsafe_allow_html=True)
        st.markdown("</div>", unsafe_allow_html=True)

        from io import BytesIO
        import docx
        doc = docx.Document()
        doc.add_heading('Cover Letter', 0)
        for para in cl.split('\n'):
            if para.strip():
                doc.add_paragraph(para.strip(), style='Normal')
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        st.download_button(
            label="📅 Download Cover Letter (DOCX)",
            data=buffer,
            file_name="Cover_Letter.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    if st.session_state.get('show_suggestions', False) and not generate_cover_letter_clicked:
        st.subheader("🔧 Upgrade Suggestions")
        with st.expander("⚙️ Advanced Options"):
            upgrade_projects = st.checkbox("Suggest Projects", key='upgrade_projects')
            upgrade_skills = st.checkbox("Suggest Skills", key='upgrade_skills')
        suggestions = None
        if (st.session_state.get('upgrade_projects') or st.session_state.get('upgrade_skills')) and resume_details:
            from app.modules.generate_doc import genai
            prompt = f"Given the following resume and job description, suggest improvements.\n\nResume:\n{resume_details.get('raw_text', '')}\n\nJob Description:\n{jd}"
            if st.session_state.get('upgrade_projects'):
                prompt += "\n\nSuggest 2-3 impactful projects to add or upgrade."
            if st.session_state.get('upgrade_skills'):
                prompt += "\n\nSuggest 2-3 in-demand skills to learn or improve."
            prompt += "\n\nIf any important certifications are missing, recommend them as well."
            model = genai.GenerativeModel('models/gemini-1.5-flash-latest')
            response = model.generate_content(prompt)
            suggestions = response.text if hasattr(response, 'text') else str(response)
            with st.expander("💡 Suggestions"):
                st.markdown(suggestions)

            if suggestions and st.button("🔄 Apply Suggestions to Resume"):
                full_prompt = f"Rewrite my resume for this job: {jd}\n\nCurrent resume:\n{resume_details.get('raw_text', '')}\n\nApply these suggestions: {suggestions}\n\nNote: Add suggested projects, skills, and certifications."
                improved_resume = model.generate_content(full_prompt)
                improved_resume_text = improved_resume.text if hasattr(improved_resume, 'text') else str(improved_resume)
                st.session_state['improved_resume_text'] = improved_resume_text
                st.markdown("<div class='resume-box'>", unsafe_allow_html=True)
                st.markdown(improved_resume_text.replace('\n', '<br>'), unsafe_allow_html=True)
                st.markdown("</div>", unsafe_allow_html=True)

                from io import BytesIO
                import re
                import docx

                clean_text = re.sub(r'[★*•●▪️-]+', '', improved_resume_text)
                clean_text = re.sub(r'\n+', '\n', clean_text)
                doc = docx.Document()
                doc.add_heading('Professional Resume', 0)
                for para in clean_text.split('\n'):
                    if para.strip():
                        doc.add_paragraph(para.strip(), style='Normal')
                buffer = BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                st.download_button(
                    label="📅 Download DOCX Resume",
                    data=buffer,
                    file_name="ATS_Optimized_Resume.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

# Mock Interview code
elif st.session_state["nav_page"] == "interview":
    st.title("🎤 Voice-Based Mock Interview")
    # Step 1: Upload Resume and JD
    with st.expander("📁 Step 1: Upload Resume and JD", expanded=True):
        col1, col2 = st.columns(2)
        with col1:
            uploaded_resume = st.file_uploader("📄 Upload Resume (PDF)", type=["pdf"], key="mock_resume")
        with col2:
            jd_text = st.text_area("📌 Paste Job Description", key="mock_jd", height=150)
    resume_text = ""
    if uploaded_resume:
        from app.modules.generate_doc.resume_parser import extract_resume_details_from_pdf
        parsed = extract_resume_details_from_pdf(uploaded_resume)
        resume_text = parsed.get("raw_text", "")
        with st.expander("📝 Extracted Resume Text"):
            st.text_area("Resume Text", resume_text, height=200, label_visibility="collapsed")

    # Step 2: Generate Questions
    with st.expander("⚙️ Interview Settings", expanded=True):
        num_questions = st.number_input("Select number of interview questions", min_value=1, max_value=20, value=5, step=1)

    if st.button("🧠 Generate Questions from Resume + JD", use_container_width=True) and (jd_text or resume_text):
        with st.spinner("🔍 Generating interview questions..."):
            from app.modules.voice_interview import get_questions_from_resume_and_jd
            question_list = get_questions_from_resume_and_jd(resume_text, jd_text)
            # Filter out coding questions (keep only viva questions)
            coding_keywords = ["code", "program", "implement", "write a function", "algorithm", "python function", "write a method"]
            def is_coding_question(q):
                ql = q["question"].lower()
                return any(kw in ql for kw in coding_keywords)
            all_viva_questions = [q for q in question_list if not is_coding_question(q)]
            # Use the number chosen by the user:
            viva_questions = all_viva_questions[:num_questions]
            st.session_state["questions"] = viva_questions
            st.session_state["current_question_idx"] = 0
            st.session_state["interview_complete"] = False
            # Reset answer state
            for key in ['recording_complete', 'audio_file_path', 'transcript', 'feedback', 'evaluation_complete']:
                if key in st.session_state:
                    del st.session_state[key]
            st.success(f"✅ {len(viva_questions)} viva questions generated successfully!")

    # Step 3: Proctor Mode and Sequential Questions
    if "questions" in st.session_state and st.session_state["questions"]:
        from app.modules.mock_interview.proctor_mode import proctor_mode_ui
        proctor_mode_ui()
        st.markdown("---")
        questions = st.session_state["questions"]
        idx = st.session_state.get("current_question_idx", 0)
        if idx >= len(questions):
            st.success("🎉 Interview complete! You have answered all questions.")
            # --- Generate and display holistic summary and report ---
            if "interview_report" not in st.session_state or "interview_summary" not in st.session_state:
                # Aggregate all feedbacks and transcripts
                all_feedbacks = st.session_state.get("all_feedbacks", [])
                all_transcripts = st.session_state.get("all_transcripts", [])
                # Get resume and JD for context
                resume_text = st.session_state.get("shared_resume_text", "")
                jd_text = st.session_state.get("shared_jd", "")
                # Simple analysis: strengths/weaknesses extraction (placeholder logic)
                strengths = []
                improvements = []
                for fb in all_feedbacks:
                    if "strength" in fb.lower() or "good" in fb.lower() or "well" in fb.lower():
                        strengths.append(fb)
                    if "improve" in fb.lower() or "weak" in fb.lower() or "could be better" in fb.lower():
                        improvements.append(fb)
                # Holistic summary (placeholder logic)
                summary = """
                <div class='card'>
                <h3>🔎 Overall Interview Summary</h3>
                <b>Resume-JD Fit:</b> {}
                <br><b>Overall Strengths:</b> {}
                <br><b>Key Areas to Improve:</b> {}
                <br><b>General Advice:</b> {}
                </div>
                """.format(
                    "Strong alignment with job requirements." if resume_text and jd_text and any(s in resume_text for s in jd_text.split()[:5]) else "Some gaps detected between resume and JD.",
                    ', '.join(strengths) if strengths else 'No major strengths detected.',
                    ', '.join(improvements) if improvements else 'No major weaknesses detected.',
                    "Review your answers and feedback above. Consider tailoring your resume more closely to the JD for best results."
                )
                st.session_state["interview_summary"] = summary
                # Compose detailed report
                report = """
                <div class='card'>   
                <h3>📝 Mock Interview Report</h3>
                <b>What Went Well:</b><br>
                <ul>{}</ul>
                <b>Areas for Improvement:</b><br>
                <ul>{}</ul>
                <b>Other Feedback:</b><br>
                <ul>{}</ul>
                </div>
                """.format(
                    ''.join(f'<li>{s}</li>' for s in strengths) if strengths else '<li>No major strengths detected.</li>',
                    ''.join(f'<li>{w}</li>' for w in improvements) if improvements else '<li>No major weaknesses detected.</li>',
                    ''.join(f'<li>{f}</li>' for f in all_feedbacks if f not in strengths+improvements) or '<li>Great effort! Review your answers above for more details.</li>'
                )
                st.session_state["interview_report"] = report
            st.markdown(st.session_state["interview_summary"], unsafe_allow_html=True)
            st.markdown(st.session_state["interview_report"], unsafe_allow_html=True)
            # Optionally, allow download
            if st.button("⬇️ Download Full Report as TXT"):
                txt = "Mock Interview Report\n\n"
                txt += "Overall Interview Summary\n" + (st.session_state["interview_summary"].replace('<div class=\'card\'>','').replace('</div>','').replace('<h3>','').replace('</h3>','').replace('<b>','').replace('</b>','').replace('<br>','\n').replace('<ul>','').replace('</ul>','').replace('<li>','- ').replace('</li>','\n')) + "\n"
                for i, (q, t, f) in enumerate(zip([q['question'] for q in questions], st.session_state.get('all_transcripts', []), st.session_state.get('all_feedbacks', [])), 1):
                    txt += f"Q{i}: {q}\nYour Answer: {t}\nFeedback: {f}\n\n"
                st.download_button("Download Report", txt, file_name="mock_interview_report.txt")
                # Reset all mock interview session state after download
                for key in [
                    "questions", "current_question_idx", "interview_complete", "selected_question",
                    "recording_complete", "audio_file_path", "transcript", "feedback", "evaluation_complete",
                    "all_feedbacks", "all_transcripts", "interview_report", "interview_summary"
                ]:
                    if key in st.session_state:
                        del st.session_state[key]
                if hasattr(st.session_state, 'audio_processor'):
                    st.session_state.audio_processor.clear_frames()
                st.success("Mock interview session has been reset. You can start a new interview!")

        else:
            current_q = questions[idx]["question"]
            st.session_state["selected_question"] = current_q
            st.subheader(f"Question {idx+1} of {len(questions)}")
            st.markdown(f"<div class='card'>{current_q}</div>", unsafe_allow_html=True)
            for key in ['recording_complete', 'audio_file_path', 'transcript', 'feedback', 'evaluation_complete']:
                if key not in st.session_state:
                    st.session_state[key] = None
            import mimetypes
            import tempfile

            audio_file = st.audio_input("🎙️ Record your answer")

            if audio_file:
                mime_type = audio_file.type  # e.g., 'audio/webm', 'audio/ogg', etc.
                ext = mimetypes.guess_extension(mime_type) or ".webm"

                with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as f:
                    f.write(audio_file.read())
                    audio_path = f.name

                st.audio(audio_path)
                from app.modules.voice_interview.record_audio import convert_to_proper_wav
                audio_path = convert_to_proper_wav(audio_path)
                if st.button("💡 Evaluate Answer", key="eval_audio_input"):
                    try:
                        result = generate_feedback_from_audio(audio_path, st.session_state["selected_question"])
                        st.session_state.transcript = result["transcript"]
                        st.session_state.feedback = result["feedback"]
                        st.session_state.evaluation_complete = True
                        st.rerun()
                    except Exception as e:
                        st.error(f"Error: {e}")

            # Uploaded Audio Option
            st.subheader("📁 Or Upload Your Answer")
            uploaded_audio = st.file_uploader("Upload your audio", type=["webm", "wav", "mp3"])
            if uploaded_audio:
                ext = os.path.splitext(uploaded_audio.name)[1]
                audio_path = save_uploaded_audio(uploaded_audio, extension=ext)
                st.audio(audio_path, format=f"audio/{ext[1:]}")
                st.success("✅ Audio uploaded successfully.")
                if st.button("💡 Evaluate Uploaded Answer"):
                    result = generate_feedback_from_audio(audio_path, st.session_state["selected_question"])
                    st.session_state.transcript = result["transcript"]
                    st.session_state.feedback = result["feedback"]
                    st.session_state.evaluation_complete = True
                    st.rerun()

            # Show feedback if available
            if st.session_state.get("evaluation_complete"):
                st.markdown("---")
                st.subheader("📝 Transcribed Answer:")
                st.markdown(f"<div class='card'>{st.session_state.get('transcript', '')}</div>", unsafe_allow_html=True)

                st.subheader("📊 Evaluation Feedback:")
                st.markdown(f"<div class='card'>{st.session_state.get('feedback', '')}</div>", unsafe_allow_html=True)

                col1, col2 = st.columns(2)
                with col1:
                    if st.button("🔄 Start Over", type="secondary"):
                        for key in ['recording_complete', 'audio_file_path', 'transcript', 'feedback', 'evaluation_complete']:
                            if key in st.session_state:
                                del st.session_state[key]
                        st.session_state.audio_processor.clear_frames()
                        st.rerun()
                with col2:
                    if st.button("➡️ Next Question", type="primary"):
                        # Save feedback and transcript for report
                        if "all_feedbacks" not in st.session_state:
                            st.session_state["all_feedbacks"] = []
                        if "all_transcripts" not in st.session_state:
                            st.session_state["all_transcripts"] = []
                        st.session_state["all_feedbacks"].append(st.session_state.get("feedback", ""))
                        st.session_state["all_transcripts"].append(st.session_state.get("transcript", ""))
                        st.session_state["current_question_idx"] = idx + 1
                        # Reset answer state
                        for key in ['recording_complete', 'audio_file_path', 'transcript', 'feedback', 'evaluation_complete']:
                            if key in st.session_state:
                                del st.session_state[key]
                        st.session_state.audio_processor.clear_frames()
                        st.rerun()

        # Proctor Evaluation (unchanged)
        if (
            hasattr(st.session_state, 'photos') and st.session_state.photos
            and hasattr(st.session_state, 'proctor_code_input') and st.session_state.proctor_code_input.strip()
        ):
            st.header("🤖 AI Proctor Evaluation (Gemini)")
            if st.button("Evaluate My Session with Gemini", key="proctor_gemini_eval"):
                with st.spinner("Sending data to Gemini for evaluation..."):
                    from app.modules.voice_interview.transcribe import transcribe_image
                    from app.modules.voice_interview.evaluate import evaluate_answer
                    last_photo = st.session_state.photos[-1][1]
                    face_eval = transcribe_image(last_photo)
                    coding_question = "Write a Python function that returns the factorial of a given number."
                    code_eval = evaluate_answer(coding_question, st.session_state.proctor_code_input)
                    st.success("Gemini Evaluation Complete!")
                    st.markdown(f"**Face/Presence Evaluation:** {face_eval}")
                    st.markdown(f"**Code Evaluation:** {code_eval['raw']}")
                    
    # Cheat Sheet code
elif st.session_state["nav_page"] == "cheatsheet":
    st.header("📚 Algorithm Cheat Sheet Generator")
    
    with st.expander("🧠 Auto-generate from Resume + Job Description", expanded=True):
        col1, col2 = st.columns(2)
        with col1:
            uploaded_resume = st.file_uploader("📄 Upload your resume (PDF)", type=["pdf"], key="cs_resume")
        with col2:
            jd_text = st.text_area("📝 Paste Job Description", key="cs_jd", height=150)

    extracted_text = ""
    if uploaded_resume:
        from app.modules.generate_doc.resume_parser import extract_resume_details_from_pdf
        resume_details = extract_resume_details_from_pdf(uploaded_resume)
        extracted_text = resume_details.get('raw_text', '')
        with st.expander("📃 Extracted Resume Text"):
            st.text_area("Resume Text", extracted_text, height=200, label_visibility="collapsed")

    if st.button("🚀 Generate Cheat Sheets", use_container_width=True):
        if not extracted_text.strip() and not jd_text.strip():
            st.warning("⚠️ Please upload a resume and/or paste a job description.")
        else:
            with st.spinner("🛠️ Generating cheat sheets..."):
                combined_text = f"{extracted_text}\n\n{jd_text}"
                topics = extract_topics_from_text(extracted_text, jd_text)
                st.success(f"🔍 Detected Topics: {', '.join(topics)}")
                for topic in topics:
                    sheet = generate_cheat_sheet(topic, jd_text)
                    with st.expander(f"📌 {topic}", expanded=False):
                        st.markdown(sheet, unsafe_allow_html=True)

    st.markdown("---")
    st.subheader("✍️ Generate by Custom Topic")
    custom_topic = st.text_input("Enter a specific algorithm topic (e.g., Binary Search Trees)")
    if st.button("📄 Generate Custom Cheat Sheet", use_container_width=True):
        with st.spinner("🛠️ Generating cheat sheet..."):
            sheet = generate_cheat_sheet(custom_topic, jd_text)
            with st.expander(f"📌 Cheat Sheet: {custom_topic}", expanded=True):
                st.markdown(sheet, unsafe_allow_html=True)

# Automated Job Hunter code
elif st.session_state["nav_page"] == "job_hunter":
    st.header("🤖 Automated Job Hunter")
    st.markdown("""
    <div class='card'>
    1. Upload or paste your resume (or use previous)<br>
    2. Paste job preferences/keywords (e.g. Data Scientist, Remote, Python)<br>
    3. Choose job source(s)<br>
    4. Click 'Find & Auto-Apply'<br>
    5. See matched jobs, generated CV/cover letter, and application status
    </div>
    """, unsafe_allow_html=True)
    
    with st.expander("📄 Resume & Preferences", expanded=True):
        uploaded_resume = st.file_uploader("📄 Upload your resume (PDF, optional)", type=["pdf"], key="auto_resume_upload")
        if 'auto_resume' not in st.session_state:
            st.session_state['auto_resume'] = ''
        if uploaded_resume:
            from app.modules.generate_doc.resume_parser import extract_resume_details_from_pdf
            resume_details = extract_resume_details_from_pdf(uploaded_resume)
            st.session_state['auto_resume'] = resume_details.get('raw_text', '')
        # Use shared resume if available and nothing uploaded
        if not uploaded_resume and st.session_state.get('shared_resume_text'):
            st.session_state['auto_resume'] = st.session_state['shared_resume_text']
        auto_resume = st.text_area("📄 Paste Resume (optional)", value=st.session_state['auto_resume'], key='auto_resume_input', height=150)
        job_keywords = st.text_area("🔎 Job Preferences/Keywords", key='auto_job_keywords', height=80)
        # Use shared JD if available and nothing entered
        if not job_keywords and st.session_state.get('shared_jd'):
            job_keywords = st.session_state['shared_jd']
        use_prev = st.button("Use Previous Resume", key='auto_use_prev', use_container_width=True)
        if use_prev:
            st.session_state['auto_resume'] = st.session_state.get('improved_resume_text') or st.session_state.get('resume', '')
            auto_resume = st.session_state['auto_resume']
        
        st.markdown("**Select job sources:**")
        col1, col2 = st.columns(2)
        with col1:
            location = st.text_input("🌍 Location (city or leave blank for India)", value="")
            use_remotive = st.checkbox("Remotive", value=True, key='use_remotive', disabled=bool(location.strip()))
        with col2:
            use_adzuna = st.checkbox("Adzuna", value=True, key='use_adzuna')

    if st.button("Find & Auto-Apply", use_container_width=True):
        with st.spinner("🔍 Searching for jobs and preparing applications..."):
            jobs = []
            # If location is given, use Adzuna only and pass location, else use Remotive (India) and Adzuna (IN)
            if not location.strip() and st.session_state.get('use_remotive'):
                jobs += fetch_jobs_remotive(job_keywords, limit=10, location='India')
            if st.session_state.get('use_adzuna'):
                adzuna_loc = location.strip() if location.strip() else 'IN'
                jobs += fetch_jobs_adzuna(job_keywords, location=adzuna_loc, limit=10)
            if not jobs:
                st.warning("No jobs found. Try different keywords or sources.")
            else:
                st.success(f"Found {len(jobs)} jobs matching your criteria.")
                from app.modules.generate_doc import genai
                model = genai.GenerativeModel('models/gemini-1.5-flash-latest')
                import re
                for i, job in enumerate(jobs, 1):
                    job_desc_clean = re.sub(r'<[^>]+>', '', job['desc'])
                    with st.expander(f"**{i}. {job['title']} at {job['company']} ({job['location']})**", expanded=False):
                        st.markdown(f"**🔗 [Job Link]({job['url']})**")
                        st.markdown(f"<div class='card'>{job_desc_clean[:500]}...</div>", unsafe_allow_html=True)
                        # --- Suggest improvements for this job ---
                        suggest_prompt = f"Given the following resume and job description, suggest improvements.\n\nResume:\n{auto_resume}\n\nJob Description:\n{job_desc_clean}"
                        suggest_prompt += "\n\nSuggest 2-3 impactful projects to add or upgrade."
                        suggest_prompt += "\n\nSuggest 2-3 in-demand skills to learn or improve."
                        suggest_prompt += "\n\nIf any important certifications are missing, recommend them as well."
                        suggestions_resp = model.generate_content(suggest_prompt)
                        suggestions = suggestions_resp.text if hasattr(suggestions_resp, 'text') else str(suggestions_resp)
                        # --- Apply suggestions to resume ---
                        apply_prompt = f"Rewrite my resume for this job: {job_desc_clean}\n\nCurrent resume:\n{auto_resume}\n\nApply these suggestions: {suggestions}\n\nNote: Add suggested projects, skills, and certifications."
                        improved_resume_resp = model.generate_content(apply_prompt)
                        improved_resume_text = improved_resume_resp.text if hasattr(improved_resume_resp, 'text') else str(improved_resume_resp)
                        # --- Clean up resume text ---
                        clean_text = re.sub(r'[★*•●▪️-]+', '', improved_resume_text)
                        clean_text = re.sub(r'\n+', '\n', clean_text)
                        st.markdown("**Tailored Resume:**")
                        st.markdown(f"<div class='resume-box'>{clean_text}</div>", unsafe_allow_html=True)
                        # --- Generate tailored cover letter ---
                        cover_letter_prompt = f"Write a professional cover letter for the following job description using this resume.\n\nJob Description:\n{job_desc_clean}\n\nResume:\n{clean_text}"
                        cl_response = model.generate_content(cover_letter_prompt)
                        cl = cl_response.text if hasattr(cl_response, 'text') else str(cl_response)
                        st.markdown("**Tailored Cover Letter:**")
                        st.markdown(f"<div class='resume-box'>{cl}</div>", unsafe_allow_html=True)
                        st.success("✅ [Demo] Application submitted! (Click job title to view/apply manually)")

st.markdown("---")
st.markdown("<center><small>🚀 Built with ❤️ using Streamlit | AI Career Companion</small></center>", unsafe_allow_html=True)